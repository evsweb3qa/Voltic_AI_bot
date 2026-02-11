# document_uploader.py
# Модуль для загрузки и обработки документов в RAG систему

import asyncio
import hashlib
import logging
from io import BytesIO
from typing import List, Dict, Any, Optional
import PyPDF2
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)


class DocumentUploader:
    """Загрузчик и обработчик документов для RAG"""

    def __init__(self, db, embedding_service):
        self.db = db
        self.embedding_service = embedding_service
        # Разделитель текста на чанки
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,  # Размер чанка в символах
            chunk_overlap=200,  # Перекрытие между чанками
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    async def process_file(self, file_bytes: bytes, filename: str, user_id: int) -> Dict[str, Any]:
        """
        Обрабатывает файл любого поддерживаемого формата.
        Поддерживаемые форматы: PDF, TXT, MD, DOCX
        """
        # Определяем тип файла по расширению
        ext = filename.lower().split('.')[-1] if '.' in filename else ''

        if ext == 'pdf':
            return await self.process_pdf(file_bytes, filename, user_id)
        elif ext in ['txt', 'md', 'text']:
            return await self.process_text(file_bytes, filename, user_id)
        elif ext in ['docx', 'doc']:
            return await self.process_docx(file_bytes, filename, user_id)
        else:
            return {
                'success': False,
                'error': f'Неподдерживаемый формат файла: .{ext}. Используйте PDF, TXT, MD или DOCX.'
            }

    async def process_pdf(self, file_bytes: bytes, filename: str, user_id: int) -> Dict[str, Any]:
        """Обрабатывает PDF файл"""
        try:
            # 1. Вычисляем хеш файла для проверки дубликатов
            file_hash = hashlib.md5(file_bytes).hexdigest()

            # 2. Проверяем, не загружался ли уже этот файл
            async with self.db.pool.acquire() as conn:
                existing = await conn.fetchval(
                    "SELECT id FROM rag_documents WHERE file_hash = $1",
                    file_hash
                )
                if existing:
                    return {
                        'success': False,
                        'error': 'Этот документ уже загружен'
                    }

            # 3. Извлекаем текст из PDF
            text = await self._extract_pdf_text(file_bytes)

            if not text or len(text.strip()) < 10:
                return {
                    'success': False,
                    'error': 'Не удалось извлечь текст из PDF (файл пустой или защищён)'
                }

            # 4. Обрабатываем текст и сохраняем
            return await self._process_and_save(text, filename, file_hash, user_id)

        except Exception as e:
            logger.error(f"Ошибка обработки PDF: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    async def process_text(self, file_bytes: bytes, filename: str, user_id: int) -> Dict[str, Any]:
        """Обрабатывает текстовый файл (TXT, MD)"""
        try:
            # 1. Вычисляем хеш файла
            file_hash = hashlib.md5(file_bytes).hexdigest()

            # 2. Проверяем дубликаты
            async with self.db.pool.acquire() as conn:
                existing = await conn.fetchval(
                    "SELECT id FROM rag_documents WHERE file_hash = $1",
                    file_hash
                )
                if existing:
                    return {
                        'success': False,
                        'error': 'Этот документ уже загружен'
                    }

            # 3. Декодируем текст (пробуем разные кодировки)
            text = None
            for encoding in ['utf-8', 'cp1251', 'latin-1']:
                try:
                    text = file_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            # Удаляем null bytes (0x00) - они недопустимы в PostgreSQL
            if text:
                text = text.replace('\x00', '')

            if not text or len(text.strip()) < 10:
                return {
                    'success': False,
                    'error': 'Не удалось прочитать текст из файла'
                }

            # 4. Обрабатываем и сохраняем
            return await self._process_and_save(text, filename, file_hash, user_id)

        except Exception as e:
            logger.error(f"Ошибка обработки текстового файла: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    async def process_docx(self, file_bytes: bytes, filename: str, user_id: int) -> Dict[str, Any]:
        """Обрабатывает DOCX файл (Word документ)"""
        try:
            # 1. Вычисляем хеш файла
            file_hash = hashlib.md5(file_bytes).hexdigest()

            # 2. Проверяем дубликаты
            async with self.db.pool.acquire() as conn:
                existing = await conn.fetchval(
                    "SELECT id FROM rag_documents WHERE file_hash = $1",
                    file_hash
                )
                if existing:
                    return {
                        'success': False,
                        'error': 'Этот документ уже загружен'
                    }

            # 3. Извлекаем текст из DOCX
            text = await self._extract_docx_text(file_bytes)

            if not text or len(text.strip()) < 10:
                return {
                    'success': False,
                    'error': 'Не удалось извлечь текст из DOCX (файл пустой или повреждён)'
                }

            # 4. Обрабатываем и сохраняем
            return await self._process_and_save(text, filename, file_hash, user_id)

        except Exception as e:
            logger.error(f"Ошибка обработки DOCX файла: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    async def _extract_docx_text(self, file_bytes: bytes) -> str:
        """Извлекает текст из DOCX файла"""
        try:
            docx_file = BytesIO(file_bytes)
            doc = docx.Document(docx_file)

            text_parts = []
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            # Объединяем и удаляем null bytes
            full_text = "\n\n".join(text_parts)
            return full_text.replace('\x00', '')

        except Exception as e:
            logger.error(f"Ошибка извлечения текста из DOCX: {e}")
            return ""

    async def _process_and_save(self, text: str, filename: str,
                                file_hash: str, user_id: int) -> Dict[str, Any]:
        """Разбивает текст на чанки, создаёт эмбеддинги и сохраняет в БД"""
        try:
            # 1. Разбиваем текст на чанки
            chunks = self.text_splitter.split_text(text)

            if not chunks:
                return {
                    'success': False,
                    'error': 'Не удалось разбить документ на части'
                }

            # 2. Добавляем документ в БД
            doc_id = await self.db.add_document(filename, file_hash, user_id)

            # 3. Обрабатываем каждый чанк
            processed_chunks = 0
            failed_chunks = 0
            for i, chunk_text in enumerate(chunks):
                # Создаём эмбеддинг для чанка
                embedding = await self.embedding_service.create_embedding(chunk_text)

                # Пропускаем чанк если эмбеддинг не создался
                if embedding is None:
                    logger.warning(f"Пропущен чанк {i} - ошибка создания эмбеддинга")
                    failed_chunks += 1
                    continue

                # Сохраняем чанк в БД
                await self.db.add_chunk(
                    document_id=doc_id,
                    chunk_index=i,
                    content=chunk_text,
                    embedding=embedding,
                    metadata={
                        'filename': filename,
                        'chunk_index': i,
                        'total_chunks': len(chunks),
                        'uploaded_by': user_id
                    }
                )
                processed_chunks += 1

            # 4. Обновляем счётчик чанков в документе
            async with self.db.pool.acquire() as conn:
                await conn.execute("""
                    UPDATE rag_documents 
                    SET total_chunks = $1 
                    WHERE id = $2
                """, processed_chunks, doc_id)

            logger.info(f"Документ {filename} обработан: {processed_chunks} чанков")

            return {
                'success': True,
                'document_id': doc_id,
                'filename': filename,
                'chunks_created': processed_chunks,
                'total_text_length': len(text)
            }

        except Exception as e:
            logger.error(f"Ошибка сохранения документа: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    async def _extract_pdf_text(self, file_bytes: bytes) -> str:
        """Извлекает текст из PDF файла"""
        try:
            pdf_file = BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text)

            # Объединяем текст и удаляем null bytes
            full_text = "\n\n".join(text_parts)
            return full_text.replace('\x00', '')

        except Exception as e:
            logger.error(f"Ошибка извлечения текста из PDF: {e}")
            return ""

    async def delete_document(self, document_id: int) -> bool:
        """Удаляет документ и все его чанки из БД"""
        try:
            return await self.db.delete_document(document_id)
        except Exception as e:
            logger.error(f"Ошибка удаления документа: {e}")
            return False

    async def get_documents_list(self) -> List[Dict]:
        """Возвращает список всех загруженных документов"""
        try:
            return await self.db.get_all_documents()
        except Exception as e:
            logger.error(f"Ошибка получения списка документов: {e}")
            return []


# Глобальный экземпляр (инициализируется при запуске)
document_uploader: Optional[DocumentUploader] = None